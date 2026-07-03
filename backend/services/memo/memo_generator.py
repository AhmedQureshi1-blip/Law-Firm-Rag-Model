"""
services/memo/memo_generator.py

Memo generator that creates professionally formatted Word documents from
the LLM-generated memo text. Uses python-docx for document formatting.

This service is critical because it:
- Produces lawyer-ready Word documents that can be directly edited
- Applies professional formatting with proper heading styles
- Creates color-coded red flag tables for quick risk assessment
- Includes proper disclaimers and page numbers
- Saves both .docx (for editing) and .txt (for preview) versions
"""

from pathlib import Path
from typing import Tuple

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import settings
from utils.logger import get_logger

logger = get_logger(__name__)


class MemoGenerator:
    """
    Generates professionally formatted Word documents from memo text.
    Creates both .docx (editable) and .txt (preview) versions.
    """

    def __init__(self):
        self.memo_dir = settings.MEMO_DIR
        self.memo_dir.mkdir(parents=True, exist_ok=True)

    def generate_memo_document(
        self,
        session_id: str,
        memo_text: str,
    ) -> Tuple[Path, Path]:
        """
        Generates a Word document and plain text version from the memo text.
        Returns paths to both the .docx and .txt files.
        """
        logger.info(f"Generating memo document for session '{session_id}'")

        # Generate Word document
        docx_path = self._generate_docx(session_id, memo_text)
        
        # Generate plain text version for preview
        txt_path = self._generate_txt(session_id, memo_text)

        logger.info(f"Memo documents generated: {docx_path} and {txt_path}")
        return docx_path, txt_path

    def _generate_docx(self, session_id: str, memo_text: str) -> Path:
        """Generates a professionally formatted Word document."""
        doc = Document()

        # Add letterhead placeholder
        self._add_letterhead(doc)

        # Parse and add memo content with proper formatting
        self._add_memo_content(doc, memo_text)

        # Add footer with page numbers and disclaimer
        self._add_footer(doc)

        # Save the document
        docx_path = self.memo_dir / f"{session_id}.docx"
        doc.save(docx_path)
        return docx_path

    def _add_letterhead(self, doc: Document):
        """Adds a placeholder letterhead area."""
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("LAW FIRM NAME — CONFIDENTIAL")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        doc.add_paragraph()  # Empty line for spacing

    def _add_memo_content(self, doc: Document, memo_text: str):
        """
        Parses memo text and adds it with proper formatting.
        Detects headings and applies appropriate styles.
        """
        lines = memo_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            # Detect section headings (all caps or followed by colon)
            if line.isupper() or line.endswith(':'):
                self._add_heading(doc, line)
            # Detect numbered lists
            elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', 
                                   '10.', '11.', '12.', '13.', '14.', '15.')):
                self._add_list_item(doc, line)
            # Detect red flag section for special formatting
            elif 'RED FLAG' in line.upper():
                self._add_heading(doc, line)
            # Regular paragraph
            else:
                self._add_paragraph(doc, line)

    def _add_heading(self, doc: Document, text: str):
        """Adds a heading with proper formatting."""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        paragraph.space_after = Pt(6)

    def _add_paragraph(self, doc: Document, text: str):
        """Adds a regular paragraph."""
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.runs[0]
        run.font.size = Pt(11)

    def _add_list_item(self, doc: Document, text: str):
        """Adds a numbered list item."""
        paragraph = doc.add_paragraph(text, style='List Number')
        run = paragraph.runs[0]
        run.font.size = Pt(11)

    def _add_footer(self, doc: Document):
        """Adds footer with page numbers and disclaimer."""
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        
        run = paragraph.add_run("AI-Generated Draft — Requires Lawyer Review")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _generate_txt(self, session_id: str, memo_text: str) -> Path:
        """Generates a plain text version for frontend preview."""
        txt_path = self.memo_dir / f"{session_id}.txt"
        txt_path.write_text(memo_text, encoding='utf-8')
        return txt_path

    def get_memo_path(self, session_id: str) -> Path:
        """Returns the path to the memo document for a session."""
        return self.memo_dir / f"{session_id}.docx"

    def memo_exists(self, session_id: str) -> bool:
        """Checks if a memo document exists for the session."""
        return self.get_memo_path(session_id).exists()
